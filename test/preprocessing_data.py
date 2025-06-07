
from docx import Document
from PIL import Image
import io
import os
from src.config.cloudinary import upload_image
def extract_chunks_from_docx_fixed(path_to_docx):
    doc = Document(path_to_docx)
    chunks = []
    current_text = []
    current_images = []
    image_counter = 0

    for para in doc.paragraphs:
        para_has_image = False
        for run in para.runs:
            # Dùng local-name để tìm ảnh
            pics = run._element.xpath(".//*[local-name()='pic']")
            if pics:
                para_has_image = True
                break
        
        if para_has_image:
            # Lưu chunk text trước đó
            if current_text or current_images:
                chunks.append({
                    "content": "\n".join(current_text).strip(),
                    "image_url": current_images.copy()
                })
                current_text = []
                current_images = []

            image_counter += 1
            url = "url" + str(image_counter)
            chunks.append({
                "content": "",
                "image_url": [url]
            })
        else:
            current_text.append(para.text)
    
    # Thêm chunk cuối cùng
    if current_text or current_images:
        chunks.append({
            "content": "\n".join(current_text).strip(),
            "image_url": current_images.copy()
        })

    return chunks



def extract_chunks_from_docx_fixed_2(path_to_docx):
    doc = Document(path_to_docx)
    chunks = []
    current_text = []
    current_images = []
    image_counter = 0

    # Create directory for saving images if it doesn't exist
    os.makedirs("extracted_images", exist_ok=True)

    for para in doc.paragraphs:
        para_has_image = False
        for run in para.runs:
            # Find images in the paragraph
            pics = run._element.xpath(".//*[local-name()='pic']")
            if pics:
                para_has_image = True
                # Extract and save image
                for pic in pics:
                    # Get the blip element which contains the image reference
                    blip = pic.find(
                        ".//a:blip",
                        {"a": "http://schemas.openxmlformats.org/drawingml/2006/main"},
                    )
                    if blip is not None:
                        embed = blip.get(
                            "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                        )
                        if embed:
                            image_part = doc.part.related_parts[embed]
                            image_data = image_part.blob
                            # Save image temporarily
                            image_counter += 1
                            image_path = f"extracted_images/image_{image_counter}.png"
                            with open(image_path, "wb") as f:
                                f.write(image_data)

                            # Upload to Cloudinary
                            upload_result = upload_image(
                                file_path=image_path,
                                folder="robokki_images",
                                public_id=f"image_{image_counter}",
                            )

                            # Get the secure URL from Cloudinary
                            cloudinary_url = upload_result["secure_url"]
                            
                            # Add image URL to both current and previous chunks if they exist
                            if chunks:
                                chunks[-1]["image_url"].append(cloudinary_url)
                            current_images.append(cloudinary_url)

                            # Clean up temporary file
                            os.remove(image_path)
                break

        if para_has_image:
            # Save previous text chunk
            if current_text or current_images:
                chunks.append(
                    {
                        "content": "\n".join(current_text).strip(),
                        "image_url": current_images.copy(),
                    }
                )
                current_text = []
                current_images = []
        else:
            current_text.append(para.text)

    # Add final chunk
    if current_text or current_images:
        chunks.append(
            {
                "content": "\n".join(current_text).strip(),
                "image_url": current_images.copy(),
            }
        )

    # Clean up the temporary directory
    if os.path.exists("extracted_images"):
        os.rmdir("extracted_images")

    return chunks


# Replace with your actual Word file path
path = "./Data/Cẩm nang copy 2.docx"
result_chunks = extract_chunks_from_docx_fixed(path)

from pprint import pprint

pprint(result_chunks)

# Thay bằng đường dẫn file Word thực tế của bạn
path = "./Data/Cẩm nang copy 2.docx"
result_chunks = extract_chunks_from_docx_fixed(path)

from pprint import pprint
pprint(result_chunks)
